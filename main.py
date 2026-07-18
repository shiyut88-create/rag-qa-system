# ========== 【必须放在最顶部】全局编码修复 ==========
import os
import sys
import tempfile

os.environ["PYTHONIOENCODING"] = "utf-8"
os.environ["PYTHONLEGACYWINDOWSSTDIO"] = "utf-8"
os.environ["HTTPX_DEFAULT_ENCODING"] = "utf-8"
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"

if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")
if sys.stderr.encoding != "utf-8":
    sys.stderr.reconfigure(encoding="utf-8")

# ========== 正式导入 ==========
import streamlit as st
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
from openai import OpenAI
import docx

# ── 页面设置 ──────────────────────────────────────────────────
st.set_page_config(page_title="RAG 文档问答", page_icon="📚")
st.title("📚 RAG 文档问答系统")
st.caption("基于 LangChain + Chroma + DeepSeek 构建")

# ── 加载 Embedding 模型（只加载一次）─────────────────────────
@st.cache_resource
def load_embedding_model():
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    )

@st.cache_resource
def load_client():
    return OpenAI(
        api_key="sk-69392588f23b491081d42b6a98f8326f",
        base_url="https://api.deepseek.com/v1"
    )

embedding_model = load_embedding_model()
client = load_client()

# ── 文件解析函数 ──────────────────────────────────────────────
def parse_file(uploaded_file):
    filename = uploaded_file.name
    suffix = os.path.splitext(filename)[1].lower()

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    docs = []

    if suffix == ".pdf":
        loader = PyMuPDFLoader(tmp_path)
        docs = loader.load()

    elif suffix == ".docx":
        doc = docx.Document(tmp_path)
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        docs = [Document(page_content=full_text, metadata={"source": filename, "page": 1})]

    elif suffix == ".txt":
        with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
            full_text = f.read()
        docs = [Document(page_content=full_text, metadata={"source": filename, "page": 1})]

    os.unlink(tmp_path)
    return docs

# ── 侧边栏 ────────────────────────────────────────────────────
with st.sidebar:
    st.header("📂 上传文档")
    uploaded_files = st.file_uploader(
        "支持 PDF、Word、TXT",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True
    )

    if uploaded_files:
        if st.button("🔄 构建知识库"):
            with st.spinner("正在解析文件并构建知识库..."):
                all_documents = []
                for uploaded_file in uploaded_files:
                    docs = parse_file(uploaded_file)
                    all_documents.extend(docs)
                    st.write(f"✔ 已解析：{uploaded_file.name}")

                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=500,
                    chunk_overlap=50
                )
                chunks = splitter.split_documents(all_documents)

                vectorstore = Chroma.from_documents(
                    documents=chunks,
                    embedding=embedding_model,
                    persist_directory="chroma_db"
                )
                st.session_state.vectorstore = vectorstore
                st.success(f"✅ 构建完成，共 {len(chunks)} 个块")

    st.divider()
    st.header("⚙️ 检索设置")
    k_value = st.slider(
        "检索块数量",
        min_value=1,
        max_value=6,
        value=3,
        help="数值越大参考内容越多，但回答速度稍慢"
    )
    st.session_state.k_value = k_value

    st.divider()
    if st.button("🗑️ 清空对话"):
        st.session_state.messages = []
        st.rerun()

# ── 问答函数（支持多轮对话）──────────────────────────────────
def ask(question, vectorstore, chat_history):
    k = st.session_state.get("k_value", 3)
    results = vectorstore.similarity_search(question, k=k)
    context = "\n\n".join([doc.page_content for doc in results])

    sources = []
    for doc in results:
        meta = doc.metadata
        filename = os.path.basename(meta.get("source", "未知文件"))
        page = meta.get("page", 0) + 1
        sources.append(f"📄 {filename}  第 {page} 页")
    sources = list(dict.fromkeys(sources))

    # 系统 prompt
    system_prompt = f"""你是一个文档助手，请根据以下参考内容回答用户的问题。
如果内容中没有相关信息，请说"文档中未找到相关内容"。

参考内容：
{context}
"""
    # 构建完整对话历史
    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": question})

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=messages
    )
    answer = response.choices[0].message.content
    return answer, sources

# ── 聊天界面 ──────────────────────────────────────────────────
if "messages" not in st.session_state:
    st.session_state.messages = []

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])
        if "sources" in msg:
            with st.expander("📌 来源"):
                for s in msg["sources"]:
                    st.write(s)

if "vectorstore" not in st.session_state:
    st.info("👈 请先在左侧上传文件并点击构建知识库")
else:
    if question := st.chat_input("请输入你的问题..."):
        st.session_state.messages.append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.write(question)

        with st.chat_message("assistant"):
            with st.spinner("正在检索文档..."):
                answer, sources = ask(
                    question,
                    st.session_state.vectorstore,
                    st.session_state.messages[:-1]  # 传入除最新问题外的历史
                )
            st.write(answer)
            with st.expander("📌 来源"):
                for s in sources:
                    st.write(s)

        st.session_state.messages.append({
            "role": "assistant",
            "content": answer,
            "sources": sources
        })
